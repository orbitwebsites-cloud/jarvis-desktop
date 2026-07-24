from __future__ import annotations

import ipaddress
import json
import os
import re
import shutil
import socket
import subprocess
import tempfile
import time
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote_plus, urlparse
from urllib.request import Request, urlopen


class _CdpSession:
    def __init__(self, websocket_url: str):
        try:
            import websocket
        except ImportError as exc:
            raise RuntimeError("Browser research support is not installed.") from exc
        self.socket = websocket.create_connection(
            websocket_url,
            timeout=7,
            suppress_origin=True,
            http_proxy_host=None,
            http_proxy_port=None,
        )
        self._next_id = 0

    def call(self, method: str, params: dict[str, Any] | None = None) -> dict[str, Any]:
        self._next_id += 1
        request_id = self._next_id
        self.socket.send(
            json.dumps({"id": request_id, "method": method, "params": params or {}})
        )
        deadline = time.monotonic() + 9
        while time.monotonic() < deadline:
            response = json.loads(self.socket.recv())
            if response.get("id") != request_id:
                continue
            if "error" in response:
                raise RuntimeError(str(response["error"].get("message", "Browser command failed.")))
            return response.get("result", {})
        raise TimeoutError(f"Browser command {method} timed out.")

    def evaluate(self, expression: str) -> Any:
        result = self.call(
            "Runtime.evaluate",
            {
                "expression": expression,
                "returnByValue": True,
                "awaitPromise": True,
            },
        )
        return (result.get("result") or {}).get("value")

    def close(self) -> None:
        try:
            self.socket.close()
        except Exception:
            pass


class WebResearcher:
    """Run bounded research in an isolated, visible Edge session."""

    def __init__(
        self,
        data_dir: Path,
        intelligence: Any | None = None,
        output_dir: Path | None = None,
    ):
        self.data_dir = Path(data_dir)
        self.intelligence = intelligence
        self.output_dir = output_dir

    @staticmethod
    def _edge_path() -> str:
        candidates = [
            shutil.which("msedge.exe"),
            str(Path(os.environ.get("PROGRAMFILES(X86)", "C:\\Program Files (x86)")) / "Microsoft" / "Edge" / "Application" / "msedge.exe"),
            str(Path(os.environ.get("PROGRAMFILES", "C:\\Program Files")) / "Microsoft" / "Edge" / "Application" / "msedge.exe"),
        ]
        for candidate in candidates:
            if candidate and Path(candidate).is_file():
                return candidate
        raise RuntimeError("Microsoft Edge is required for isolated web research.")

    @staticmethod
    def _free_port() -> int:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as stream:
            stream.bind(("127.0.0.1", 0))
            return int(stream.getsockname()[1])

    @staticmethod
    def _safe_public_url(value: str) -> bool:
        try:
            parsed = urlparse(value)
            if parsed.scheme not in {"http", "https"} or not parsed.hostname:
                return False
            host = parsed.hostname.lower()
            if host in {"localhost", "localhost.localdomain"} or host.endswith(".local"):
                return False
            try:
                address = ipaddress.ip_address(host)
            except ValueError:
                try:
                    addresses = {
                        item[4][0]
                        for item in socket.getaddrinfo(host, None, type=socket.SOCK_STREAM)
                    }
                except OSError:
                    return False
                return bool(addresses) and all(
                    not (
                        (address := ipaddress.ip_address(item)).is_private
                        or address.is_loopback
                        or address.is_link_local
                        or address.is_reserved
                    )
                    for item in addresses
                )
            return not (
                address.is_private
                or address.is_loopback
                or address.is_link_local
                or address.is_reserved
            )
        except ValueError:
            return False

    @staticmethod
    def _wait_for_target(port: int, timeout: float = 12) -> dict[str, Any]:
        deadline = time.monotonic() + timeout
        while time.monotonic() < deadline:
            try:
                request = Request(
                    f"http://127.0.0.1:{port}/json/list",
                    headers={"Accept": "application/json"},
                )
                with urlopen(request, timeout=1) as response:
                    targets = json.loads(response.read().decode("utf-8"))
                pages = [
                    item
                    for item in targets
                    if item.get("type") == "page" and item.get("webSocketDebuggerUrl")
                ]
                if pages:
                    return pages[0]
            except (OSError, ValueError, TimeoutError):
                pass
            time.sleep(0.2)
        raise RuntimeError("The isolated research browser did not start.")

    @staticmethod
    def _wait_for_page(session: _CdpSession, timeout: float = 6) -> None:
        deadline = time.monotonic() + timeout
        while time.monotonic() < deadline:
            if session.evaluate("document.readyState") in {"interactive", "complete"}:
                time.sleep(0.4)
                return
            time.sleep(0.15)
        raise TimeoutError("A research page took too long to load.")

    @staticmethod
    def _clean_text(value: str) -> str:
        value = value.replace("\x00", "")
        value = re.sub(r"\n{3,}", "\n\n", value)
        return value.strip()

    def _navigate(self, session: _CdpSession, url: str) -> None:
        session.call("Page.navigate", {"url": url})
        self._wait_for_page(session)

    def _collect_search_links(self, session: _CdpSession, limit: int) -> list[str]:
        expression = """
        (() => {
          const primary = [...document.querySelectorAll(
            'li.b_algo h2 a, a[data-testid="result-title-a"], .result__a, #search a[href], main a[href]'
          )];
          return [...new Set(primary.map(a => a.href).filter(Boolean))].slice(0, 40);
        })()
        """
        links = session.evaluate(expression) or []
        result = []
        for link in links:
            if (
                isinstance(link, str)
                and self._safe_public_url(link)
                and "bing.com/search" not in link
                and "duckduckgo.com/?q=" not in link
                and "google.com/search" not in link
                and link not in result
            ):
                result.append(link)
                if len(result) >= limit:
                    break
        return result

    def _extract_page(self, session: _CdpSession, url: str) -> dict[str, str]:
        self._navigate(session, url)
        payload = session.evaluate(
            """
            (() => {
              const root = document.querySelector('article, main, [role="main"]') || document.body;
              const text = (root?.innerText || document.body?.innerText || '').slice(0, 18000);
              const description = document.querySelector('meta[name="description"]')?.content || '';
              return { title: document.title || location.hostname, url: location.href, description, text };
            })()
            """
        )
        if not isinstance(payload, dict):
            raise RuntimeError("The page did not expose readable content.")
        final_url = str(payload.get("url", url))
        if not self._safe_public_url(final_url):
            raise RuntimeError("A research page redirected to a non-public address.")
        return {
            "title": self._clean_text(str(payload.get("title", "")))[:300],
            "url": final_url,
            "description": self._clean_text(str(payload.get("description", "")))[:1000],
            "text": self._clean_text(str(payload.get("text", "")))[:18000],
        }

    @staticmethod
    def _slug(value: str) -> str:
        slug = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
        return slug[:55] or "research"

    @staticmethod
    def _process_tree(process: subprocess.Popen) -> list[Any]:
        try:
            import psutil

            return [psutil.Process(process.pid), *psutil.Process(process.pid).children(recursive=True)]
        except Exception:
            return []

    @staticmethod
    def _finish_process_tree(
        process: subprocess.Popen,
        tracked_processes: list[Any],
        profile_dir: Path,
    ) -> None:
        try:
            process.wait(timeout=4)
        except subprocess.TimeoutExpired:
            process.terminate()
        for tracked in reversed(tracked_processes):
            try:
                if tracked.is_running():
                    tracked.terminate()
            except Exception:
                continue
        if tracked_processes:
            try:
                import psutil

                _, alive = psutil.wait_procs(tracked_processes, timeout=3)
                for tracked in alive:
                    try:
                        tracked.kill()
                    except Exception:
                        pass
                psutil.wait_procs(alive, timeout=2)
            except Exception:
                pass
        try:
            import psutil

            profile_marker = str(profile_dir).lower()
            profile_processes = []
            for candidate in psutil.process_iter(["cmdline"]):
                try:
                    command_line = " ".join(candidate.info.get("cmdline") or []).lower()
                    if profile_marker in command_line:
                        profile_processes.append(candidate)
                except (psutil.AccessDenied, psutil.NoSuchProcess):
                    continue
            for candidate in profile_processes:
                try:
                    candidate.terminate()
                except (psutil.AccessDenied, psutil.NoSuchProcess):
                    pass
            _, alive = psutil.wait_procs(profile_processes, timeout=3)
            for candidate in alive:
                try:
                    candidate.kill()
                except (psutil.AccessDenied, psutil.NoSuchProcess):
                    pass
            psutil.wait_procs(alive, timeout=2)
        except Exception:
            pass
        try:
            process.wait(timeout=2)
        except subprocess.TimeoutExpired:
            process.kill()
            process.wait(timeout=2)

    @staticmethod
    def _remove_profile(profile_dir: Path) -> None:
        for _ in range(20):
            shutil.rmtree(profile_dir, ignore_errors=True)
            if not profile_dir.exists():
                return
            time.sleep(0.5)

    def _write_document(
        self,
        query: str,
        summary: str,
        sources: list[dict[str, str]],
    ) -> Path:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Word document export support is not installed.") from exc
        output_dir = self.output_dir or Path.home() / "Documents" / "JARVIS Research"
        output_dir.mkdir(parents=True, exist_ok=True)
        output = output_dir / f"{datetime.now():%Y%m%d-%H%M%S}-{self._slug(query)}.docx"
        document = Document()
        document.add_heading(f"JARVIS Research: {query}", level=0)
        document.add_paragraph(f"Generated {datetime.now():%B %d, %Y at %I:%M %p}")
        document.add_heading("Summary", level=1)
        for paragraph in summary.split("\n\n"):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
        document.add_heading("Sources and extracted notes", level=1)
        for index, source in enumerate(sources, 1):
            document.add_heading(f"[{index}] {source['title']}", level=2)
            document.add_paragraph(source["url"])
            excerpt = source["description"] or source["text"][:2500]
            document.add_paragraph(excerpt)
        document.add_paragraph(
            "Generated from a bounded public-web research session. Verify important claims against the linked sources."
        )
        document.save(output)
        return output

    def research(self, query: str, max_sources: int = 4) -> dict[str, Any]:
        query = " ".join(query.split())
        if not query:
            raise ValueError("A research question is required.")
        max_sources = max(1, min(int(max_sources), 6))
        self.data_dir.mkdir(parents=True, exist_ok=True)
        profile_dir = Path(
            tempfile.mkdtemp(prefix="research-browser-", dir=str(self.data_dir))
        )
        port = self._free_port()
        search_url = "https://www.bing.com/search?q=" + quote_plus(query)
        process = subprocess.Popen(
            [
                self._edge_path(),
                f"--remote-debugging-port={port}",
                "--remote-allow-origins=*",
                f"--user-data-dir={profile_dir}",
                "--no-first-run",
                "--disable-default-apps",
                "--disable-background-mode",
                "--new-window",
                search_url,
            ],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        session: _CdpSession | None = None
        sources: list[dict[str, str]] = []
        try:
            target = self._wait_for_target(port)
            session = _CdpSession(target["webSocketDebuggerUrl"])
            session.call("Page.enable")
            session.call("Runtime.enable")
            self._wait_for_page(session)
            search_pages = [
                search_url,
                "https://duckduckgo.com/?q=" + quote_plus(query),
                "https://www.google.com/search?q=" + quote_plus(query),
            ]
            attempted_links: set[str] = set()
            attempt_limit = max(5, max_sources * 2)
            research_deadline = time.monotonic() + 80
            for index, engine_url in enumerate(search_pages):
                if time.monotonic() >= research_deadline:
                    break
                if index:
                    self._navigate(session, engine_url)
                links = self._collect_search_links(
                    session, min(24, max_sources * 4)
                )
                for link in links:
                    if time.monotonic() >= research_deadline:
                        break
                    if link in attempted_links:
                        continue
                    attempted_links.add(link)
                    try:
                        source = self._extract_page(session, link)
                        if len(source["text"]) >= 120:
                            sources.append(source)
                    except Exception:
                        pass
                    if len(sources) >= max_sources or len(attempted_links) >= attempt_limit:
                        break
                if len(sources) >= max_sources or len(attempted_links) >= attempt_limit:
                    break
            if not attempted_links:
                raise RuntimeError("The search pages returned no public result links.")
            if not sources:
                raise RuntimeError("The result pages did not expose readable content.")
            if self.intelligence:
                summary = self.intelligence.summarize_research(query, sources)
            else:
                summary = "\n\n".join(
                    f"[{index}] {item['title']}: {item['description'] or item['text'][:500]}"
                    for index, item in enumerate(sources, 1)
                )
            output = self._write_document(query, summary, sources)
            return {
                "query": query,
                "summary": summary,
                "path": str(output),
                "sources": [
                    {"title": item["title"], "url": item["url"]}
                    for item in sources
                ],
            }
        finally:
            tracked_processes = self._process_tree(process)
            if session:
                try:
                    session.call("Browser.close")
                except Exception:
                    pass
                session.close()
            self._finish_process_tree(process, tracked_processes, profile_dir)
            self._remove_profile(profile_dir)
